from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

def create_roleplay_docx():
    # Create a new document
    doc = Document()
    
    # Title
    title = doc.add_heading('🎭 FixMo Role-Playing Instructions', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Character Assignment
    doc.add_heading('👥 Character Assignment', level=1)
    
    characters = [
        ('Jong', 'Support Agent', 'Helps with any issues and technical problems'),
        ('Sara', 'Payment Processor', 'Handles all money transactions and escrow'),
        ('Chloe', 'Tasker (Service Provider)', 'Does the actual work'),
        ('Cheska', 'Client (Task Creator)', 'Needs work done and manages the task'),
        ('Chance', 'Quality Inspector', 'Reviews work quality and manages ratings')
    ]
    
    for name, role, description in characters:
        p = doc.add_paragraph()
        p.add_run(f'**{name}** = **{role}** - {description}').bold = True
    
    doc.add_page_break()
    
    # Complete Workflow
    doc.add_heading('📋 Complete Workflow', level=1)
    
    # Step 1
    doc.add_heading('Step 1: Client Registration & Task Creation', level=2)
    p = doc.add_paragraph()
    p.add_run('Cheska (Client): "I need someone to help me clean and organize my home office"').italic = True
    
    doc.add_heading('Cheska\'s Actions:', level=3)
    actions = [
        'Open FixMo app',
        'Register as a new user',
        'Complete profile setup',
        'Navigate to "Post Task"',
        'Fill out task details:',
        '  - Title: "Home Office Organization & Cleaning"',
        '  - Description: "Need help cleaning desk, organizing files, and setting up a productive workspace"',
        '  - Budget: ₱6,000',
        '  - Location: Home address',
        '  - Date: This Friday',
        '  - Duration: 3 hours',
        'Post the task',
        'Wait for tasker applications'
    ]
    
    for action in actions:
        doc.add_paragraph(action, style='List Bullet')
    
    # Step 2
    doc.add_heading('Step 2: Tasker Discovery & Application', level=2)
    p = doc.add_paragraph()
    p.add_run('Chloe (Tasker): "I love organizing spaces and I\'m available Friday"').italic = True
    
    doc.add_heading('Chloe\'s Actions:', level=3)
    actions = [
        'Open FixMo app',
        'Register as a tasker',
        'Complete verification process (selfie + ID)',
        'Browse available tasks',
        'Find Cheska\'s office task',
        'Apply for the task with:',
        '  - Proposal: "I specialize in home organization and can create a beautiful, functional workspace"',
        '  - Price: ₱5,500 (within budget)',
        '  - Availability: Friday 2 PM - 5 PM',
        '  - References: Previous office organization projects',
        'Send application'
    ]
    
    for action in actions:
        doc.add_paragraph(action, style='List Bullet')
    
    # Step 3
    doc.add_heading('Step 3: Client Review & Selection', level=2)
    p = doc.add_paragraph()
    p.add_run('Cheska (Client): "Let me review the applications and choose the best tasker"').italic = True
    
    doc.add_heading('Cheska\'s Actions:', level=3)
    actions = [
        'Check notifications for new applications',
        'Review Chloe\'s profile and application',
        'Check her ratings and reviews',
        'View her previous work examples',
        'Send message: "Hi Chloe, your portfolio looks amazing! Can you start at 2 PM?"',
        'Chloe responds: "Absolutely! I\'ll bring some organizing supplies and cleaning products."',
        'Accept Chloe\'s application',
        'Confirm task details'
    ]
    
    for action in actions:
        doc.add_paragraph(action, style='List Bullet')
    
    # Step 4
    doc.add_heading('Step 4: Payment Setup & Escrow', level=2)
    p = doc.add_paragraph()
    p.add_run('Sara (Payment Processor): "Let\'s secure the payment for this task"').italic = True
    
    doc.add_heading('Sara\'s Actions:', level=3)
    actions = [
        'Guide Cheska through payment setup',
        'Cheska selects payment method (GCash)',
        'Enter payment details',
        'Confirm ₱5,500 payment',
        'Payment goes to escrow (held safely)',
        'Send confirmation to both Cheska and Chloe',
        'Task is now "Paid & Scheduled"'
    ]
    
    for action in actions:
        doc.add_paragraph(action, style='List Bullet')
    
    # Step 5
    doc.add_heading('Step 5: Task Execution & Communication', level=2)
    p = doc.add_paragraph()
    p.add_run('Chloe (Tasker): "I\'m on my way to organize the home office"').italic = True
    
    doc.add_heading('Chloe\'s Actions:', level=3)
    actions = [
        'Receive task confirmation',
        'Check task details and location',
        'Send "On my way" message',
        'Arrive at location',
        'Start the task',
        'Send progress updates with photos',
        'Complete the organization work',
        'Send "Task completed" notification'
    ]
    
    for action in actions:
        doc.add_paragraph(action, style='List Bullet')
    
    doc.add_heading('Cheska\'s Actions:', level=3)
    actions = [
        'Receive arrival notification',
        'Meet Chloe at the office',
        'Monitor progress through app',
        'Receive progress photos',
        'Approve completed work',
        'Release payment from escrow'
    ]
    
    for action in actions:
        doc.add_paragraph(action, style='List Bullet')
    
    # Step 6
    doc.add_heading('Step 6: Quality Review & Rating', level=2)
    p = doc.add_paragraph()
    p.add_run('Chance (Quality Inspector): "Let\'s review the completed work"').italic = True
    
    doc.add_heading('Chance\'s Actions:', level=3)
    actions = [
        'Review completed task photos',
        'Check if work meets requirements',
        'Verify task completion criteria',
        'Guide Cheska through rating process',
        'Cheska rates Chloe: 5 stars',
        'Cheska leaves review: "Fantastic work! My office is now beautiful and functional"',
        'Chloe rates Cheska: 5 stars',
        'Chloe leaves review: "Great client, very clear about what she wanted"'
    ]
    
    for action in actions:
        doc.add_paragraph(action, style='List Bullet')
    
    # Step 7
    doc.add_heading('Step 7: Support & Issue Resolution', level=2)
    p = doc.add_paragraph()
    p.add_run('Jong (Support Agent): "I\'m here to help with any issues"').italic = True
    
    doc.add_heading('Jong\'s Actions:', level=3)
    actions = [
        'Monitor the entire process',
        'Be ready to help with any problems',
        'If issues arise:',
        '  - Help with payment problems',
        '  - Resolve communication issues',
        '  - Handle disputes',
        '  - Provide technical support',
        'Ensure smooth experience'
    ]
    
    for action in actions:
        doc.add_paragraph(action, style='List Bullet')
    
    doc.add_page_break()
    
    # What Each Person Tests
    doc.add_heading('🎯 What Each Person Tests', level=1)
    
    test_roles = [
        ('Cheska (Client)', [
            'User registration and onboarding',
            'Task creation and posting',
            'Budget setting in Philippine Pesos',
            'Tasker selection and communication',
            'Task monitoring and completion',
            'Work approval process',
            'Rating and review system'
        ]),
        ('Chloe (Tasker)', [
            'Tasker registration and verification',
            'Task discovery and application',
            'Client communication',
            'Task execution tools',
            'Progress tracking',
            'Payment receipt in Philippine Pesos'
        ]),
        ('Sara (Payment)', [
            'Philippine payment method setup (GCash, PayMaya, GoTyme, Cash)',
            'Escrow system in Philippine Pesos',
            'Transaction processing',
            'Payment security',
            'Refund handling',
            'Financial reporting'
        ]),
        ('Chance (Quality)', [
            'Quality assessment tools',
            'Rating system',
            'Review process',
            'Dispute resolution',
            'Quality standards',
            'Feedback collection'
        ]),
        ('Jong (Support)', [
            'Help system',
            'Issue reporting',
            'Support communication',
            'Problem resolution',
            'User guidance',
            'System reliability'
        ])
    ]
    
    for role, tests in test_roles:
        doc.add_heading(f'{role} Tests:', level=2)
        for test in tests:
            doc.add_paragraph(test, style='List Bullet')
    
    doc.add_page_break()
    
    # Testing Checklist
    doc.add_heading('📱 Testing Checklist', level=1)
    
    doc.add_heading('Before Starting:', level=2)
    before_items = [
        'Everyone has the app installed',
        'Test accounts are ready',
        'Philippine payment methods are set up',
        'Communication channels are working'
    ]
    for item in before_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('During the Workflow:', level=2)
    during_items = [
        'Each step is completed successfully',
        'All features work as expected',
        'Communication flows smoothly',
        'Philippine Peso payments process correctly',
        'Quality review is thorough'
    ]
    for item in during_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('After Completion:', level=2)
    after_items = [
        'Task is successfully completed',
        'Payment is properly released in Philippine Pesos',
        'Ratings and reviews are submitted',
        'All parties are satisfied',
        'No issues or bugs encountered'
    ]
    for item in after_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # Role-Playing Tips
    doc.add_heading('🎭 Role-Playing Tips', level=1)
    tips = [
        'Stay in Character: Respond as your assigned role would',
        'Follow the Flow: Complete each step before moving to the next',
        'Document Issues: Note any problems or improvements needed',
        'Be Realistic: Use realistic Philippine Peso amounts and scenarios',
        'Have Fun: Make it enjoyable while being thorough'
    ]
    for tip in tips:
        doc.add_paragraph(tip, style='List Number')
    
    # Payment Information
    doc.add_heading('Payment Amounts:', level=2)
    payment_info = [
        'Task Budget: ₱6,000',
        'Tasker Price: ₱5,500',
        'Platform Fee: ₱275 (5%)',
        'Tasker Receives: ₱5,225',
        'Escrow Amount: ₱5,500'
    ]
    for info in payment_info:
        doc.add_paragraph(info, style='List Bullet')
    
    doc.add_heading('Philippine Payment Methods:', level=2)
    payment_methods = [
        'GCash',
        'PayMaya',
        'GoTyme',
        'Cash',
        'Bank Transfer',
        'Credit/Debit Cards'
    ]
    for method in payment_methods:
        doc.add_paragraph(method, style='List Bullet')
    
    # Save the document
    doc.save('FixMo_RolePlay_Instructions.docx')
    print("DOCX file created successfully: FixMo_RolePlay_Instructions.docx")

if __name__ == "__main__":
    create_roleplay_docx() 